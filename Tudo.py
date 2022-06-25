import mysql.connector, requests, os, win32com.client
import googlemaps
import sys
from PyQt5 import sip
from PyQt5 import QtGui
from PyQt5.QtGui import *
from docx import Document
from PyQt5.QtCore import *
from datetime import datetime
from PyQt5.QtWidgets import *
from docx.shared import Inches, Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH

class JanelaInicial(QMainWindow):
    def __init__(self):
        super().__init__()
        self.frameCadastro()
        self.frameLogin()
        self.framePerfil()
        self.carrega_janela()

    def carrega_janela(self):
        self.setGeometry(200, 60, 850, 550)
        self.setWindowTitle('TransportFlash')
        self.setStyleSheet('background-color: #f0e1f5')
        self.icone = 'minicaminhao.png'
        self.setWindowIcon(QtGui.QIcon(self.icone))

    def framePerfil(self):
        self.frame1 = QFrame(self)
        self.frame1.setGeometry(425, 0, 430, 550)
        self.frame1.setStyleSheet("QFrame {background-color: #9c32bf}")
        self.frame1.setFrameShape(QFrame.StyledPanel)
        self.frame1.setFrameShadow(QFrame.Raised)

        #imagens
        self.image_perfil = QLabel(self.frame1)
        self.image_perfil.move(80, 50)
        self.image_perfil.setPixmap(QtGui.QPixmap('''icon_perfil.png'''))

        self.image_car = QLabel(self.frame1)
        self.image_car.move(270, 430)
        self.image_car.resize(75, 45)
        self.image_car.setPixmap(QtGui.QPixmap('caminhao.png'))

        #textos
        self.label= QLabel(self.frame1)
        self.label.setText('Transport')
        self.label.move(100, 440)
        self.label.resize(110, 30)
        font = QFont()
        font.setFamily("Cascadia Mono SemiBold")
        font.setPointSize(15)
        font.setItalic(True)
        self.label.setStyleSheet('color: white')
        self.label.setFont(font)

        self.label2= QLabel(self.frame1)
        self.label2.setText('Flash')
        self.label2.move(210, 440)
        self.label2.resize(60, 30)
        self.label2.setStyleSheet('color: #edb61f')
        self.label2.setFont(font)

        self.label3= QLabel(self.frame1)
        self.label3.setText('''Somos a melhor transportadora 
              de Rio Verde-GO.''')
        self.label3.move(105, 335)
        fonte = QFont()
        fonte.setFamily('Calibri')
        fonte.setPointSize(13)
        fonte.setBold(True)
        self.label3.setStyleSheet('color: #edb61f')
        self.label3.setFont(fonte)

    def frameLogin(self):
        self.frame2 = QFrame(self)
        self.frame2.setGeometry(0, 0, 425, 550)
        self.frame2.setFrameShape(QFrame.StyledPanel)
        self.frame2.setFrameShadow(QFrame.Raised)

        #texto
        label = QLabel(self.frame2)
        label.setText('LOGIN')
        label.setGeometry(120, 50, 200, 100)
        label.setStyleSheet('color: #9c32bf')
        font = label.font()
        font.setPointSize(40)
        font.setFamily('Cascadia Mono')
        font.setItalic(True)
        label.setFont(font)

        #entradas
        estilo = '''QLineEdit{background-color: white;
                                padding-left: 10px;
                                font-size: 11pt; 
                                font-family: arial;
                                font-weight: 480;
                                color: rgba(0, 0, 0, 0.8)}'''

        self.entrada1_log = QLineEdit(self.frame2)
        self.entrada1_log.setGeometry(65, 220, 280, 40)
        self.entrada1_log.setStyleSheet(estilo)
        self.entrada1_log.setPlaceholderText('Seu CPF:')

        self.entrada2_log = QLineEdit(self.frame2)
        self.entrada2_log.setEchoMode(QLineEdit.Password)
        self.entrada2_log.setGeometry(65, 300, 280, 40)
        self.entrada2_log.setStyleSheet(estilo)
        self.entrada2_log.setPlaceholderText('Senha:')

        #botões
        estilo_botao = '''QPushButton 
                                    {background-color: #9c32bf;
                                    border-radius: 10px;
                                    color: white;}
                                    QPushButton:hover {
                                        background-color: #fff;
                                        color: #9c32bf;
                                        border: 1px solid #9c32bf;
                                    }'''
        self.botao_senha_perdida = QPushButton('Esqueceu a senha?', self.frame2)
        self.botao_senha_perdida.setGeometry(240, 350, 100, 20)
        self.botao_senha_perdida.setStyleSheet('''color: #edb61f;
                                    background-color: #f0e1f5;
                                    border: none''')

        self.botao_logar = QPushButton('Logar', self.frame2)
        self.botao_logar.setGeometry(80, 450, 80, 40)
        self.botao_logar.setStyleSheet(estilo_botao)
        self.botao_logar.setFont(QFont('Calibri', 15))
        self.botao_logar.clicked.connect(self.logar)

        self.botao_cadastrar = QPushButton('Cadastrar', self.frame2)
        self.botao_cadastrar.setGeometry(220, 450, 120, 40)
        self.botao_cadastrar.setStyleSheet(estilo_botao)
        self.botao_cadastrar.setFont(QFont('Calibri', 15))
        self.botao_cadastrar.clicked.connect(self.destruir_frameLogin)
        self.frame2.show()

    def frameCadastro(self):
        self.frame3 = QFrame(self)
        self.frame3.setGeometry(0, 0, 425, 550)
        self.frame3.setFrameShape(QFrame.StyledPanel)
        self.frame3.setFrameShadow(QFrame.Raised)

        #texto
        label = QLabel(self.frame3)
        label.setText('CADASTRAR')
        label.setGeometry(90, 20, 260, 100)
        label.setStyleSheet('color: #9c32bf')
        font = label.font()
        font.setPointSize(28)
        font.setFamily('Cascadia Mono')
        font.setItalic(True)
        label.setFont(font)

        #entradas
        self.entrada1 = QLineEdit(self.frame3)
        self.entrada1.setGeometry(70, 120, 280, 30)
        self.entrada1.setStyleSheet('''QLineEdit{background-color: white;
                                    padding-left: 10px;
                                    font-size: 11pt; 
                                    font-family: arial;
                                    font-weight: 480;
                                    color: rgba(0, 0, 0, 0.8)}''')
        self.entrada1.setPlaceholderText('Nome:')
        font = self.entrada1.font()
        font.setPointSize(14)
        font.setFamily('Calibri')
        self.entrada1.setFont(font)

        self.entrada2 = QLineEdit(self.frame3)
        self.entrada2.setGeometry(70, 170, 280, 30)
        self.entrada2.setStyleSheet('''QLineEdit{background-color: white;
                                    padding-left: 10px;
                                    font-size: 11pt; 
                                    font-family: arial;
                                    font-weight: 480;
                                    color: rgba(0, 0, 0, 0.8)}''')
        self.entrada2.setPlaceholderText('CPF:')
        font = self.entrada2.font()
        font.setPointSize(14)
        font.setFamily('Calibri')
        self.entrada2.setFont(font)

        self.entrada3 = QLineEdit(self.frame3)
        self.entrada3.setGeometry(70, 220, 280, 30)
        self.entrada3.setStyleSheet('''QLineEdit{background-color: white;
                                    padding-left: 10px;
                                    font-size: 11pt; 
                                    font-family: arial;
                                    font-weight: 480;
                                    color: rgba(0, 0, 0, 0.8)}''')
        self.entrada3.setPlaceholderText('Email:')
        font = self.entrada3.font()
        font.setPointSize(14)
        font.setFamily('Calibri')
        self.entrada3.setFont(font)

        self.entrada4 = QLineEdit(self.frame3)
        self.entrada4.setGeometry(70, 270, 280, 30)
        self.entrada4.setStyleSheet('''QLineEdit{background-color: white;
                                    padding-left: 10px;
                                    font-size: 11pt; 
                                    font-family: arial;
                                    font-weight: 480;
                                    color: rgba(0, 0, 0, 0.8)}''')
        self.entrada4.setPlaceholderText('Número de telefone:')
        font = self.entrada4.font()
        font.setPointSize(14)
        font.setFamily('Calibri')
        self.entrada4.setFont(font)

        self.entrada5 = QLineEdit(self.frame3)
        self.entrada5.setGeometry(70, 320, 280, 30)
        self.entrada5.setStyleSheet('''QLineEdit{background-color: white;
                                    padding-left: 10px;
                                    font-size: 11pt; 
                                    font-family: arial;
                                    font-weight: 480;
                                    color: rgba(0, 0, 0, 0.8)}''')
        self.entrada5.setPlaceholderText('Número do cartão:')
        font = self.entrada5.font()
        font.setPointSize(14)
        font.setFamily('Calibri')
        self.entrada5.setFont(font)

        self.entrada6 = QLineEdit(self.frame3)
        self.entrada6.setGeometry(70, 370, 280, 30)
        self.entrada6.setStyleSheet('''QLineEdit{background-color: white;
                                    padding-left: 10px;
                                    font-size: 11pt; 
                                    font-family: arial;
                                    font-weight: 480;
                                    color: rgba(0, 0, 0, 0.8)}''')
        self.entrada6.setPlaceholderText('Senha:')
        font = self.entrada6.font()
        font.setPointSize(14)
        font.setFamily('Calibri')
        self.entrada6.setFont(font)        

        #botões
        estilo_botao = '''QPushButton 
                                    {background-color: #9c32bf;
                                    border-radius: 10px;
                                    color: white;}
                                    QPushButton:hover {
                                        background-color: #fff;
                                        color: #9c32bf;
                                        border: 1px solid #9c32bf;
                                    }'''
        self.botao1 = QPushButton('Logar', self.frame3)
        self.botao1.setGeometry(80, 450, 80, 40)
        self.botao1.setStyleSheet(estilo_botao)
        self.botao1.setFont(QFont('Calibri', 15))
        self.botao1.clicked.connect(self.destruir_frameCadastro)

        self.botao2 = QPushButton('Cadastrar', self.frame3)
        self.botao2.setGeometry(220, 450, 120, 40)
        self.botao2.setStyleSheet(estilo_botao)
        self.botao2.setFont(QFont('Calibri', 15))
        self.botao2.clicked.connect(self.cadastrar)
        self.frame3.show()

    def destruir_frameLogin(self):
        sip.delete(self.frame2)
        self.frameCadastro()
        
    def destruir_frameCadastro(self):
        sip.delete(self.frame3)
        self.frameLogin()

    def conectar_no_banco(self):
        self.conexao = mysql.connector.connect(
        host = "localhost",
        user = "root",
        passwd = "",
        database = "transportadora")

        self.cursor = self.conexao.cursor()

    def desconectar_do_banco(self):
        self.conexao.close()

    def cadastrar(self):
        self.nome  = self.entrada1.text()
        self.cpf = self.entrada2.text()
        self.email = self.entrada3.text()
        self.telefone = self.entrada4.text()
        self.ddd = self.telefone[:2]
        self.numero = self.telefone[2:]
        self.cartao = self.entrada5.text()
        self.senha = self.entrada6.text()
        print(self.nome, self.cpf, self.email, self.telefone, self.cartao, self.senha)

        self.conectar_no_banco()
        self.cursor.execute(f'''INSERT INTO clientes(cpf, nome, senha, numero_do_cartao, email, ddd, numero) 
                            VALUES('{self.cpf}', '{self.nome}', '{self.senha}', '{self.cartao}', '{self.email}', '{self.ddd}', '{self.numero}');''')
        self.conexao.commit()

        self.entrada1.clear()
        self.entrada2.clear()
        self.entrada3.clear()
        self.entrada4.clear()
        self.entrada5.clear()
        self.entrada6.clear()

    def logar(self):
        self.cpf = self.entrada1_log.text()
        senha = self.entrada2_log.text()

        self.conectar_no_banco()
        self.cursor.execute(f'''SELECT senha FROM clientes WHERE cpf = '{self.cpf}';''')
        senha_user = self.cursor.fetchall()

        if senha == senha_user[0][0]:
            sip.delete(self.frame1)
            sip.delete(self.frame2)
            self.frame_BarraSuperior()
            self.frameMercadorias()
        else:
            print('senha errada')
        try:
            sip.delete(self.frame3)
        except:
            pass

    def frame_BarraSuperior(self):
        self.frame4 = QFrame(self)
        self.frame4.setGeometry(0, 0, 850, 55)
        self.frame4.setStyleSheet('''QFrame {background-color: #9c32bf}''')
        self.frame4.setFrameShape(QFrame.StyledPanel)
        self.frame4.setFrameShadow(QFrame.Raised)

        #image
        image_car = QLabel(self.frame4)
        image_car.setGeometry(180, 10, 80, 40)
        image_car.setPixmap(QtGui.QPixmap('caminhao.png'))

        #texto
        label= QLabel(self.frame4)
        label.setText('Transport')
        label.setGeometry(10, 10, 171, 31)
        font = QFont()
        font.setFamily("Cascadia Mono SemiBold")
        font.setPointSize(15)
        font.setItalic(True)
        label.setStyleSheet('color: white')
        label.setFont(font)

        label2 = QLabel(self.frame4)
        label2.setText('Flash')
        label2.setGeometry(120, 15, 71, 21)
        label2.setStyleSheet('color: #edb61f')
        label2.setFont(font)

        #botões
        botao = QPushButton(self.frame4)
        botao.setGeometry(795, 0, 55, 55)
        botao.setStyleSheet('''QPushButton{background-image: url(perfil.png);
                                background-color: #9c32bf;
                                border: 1px solid violet}
                                QPushButton:hover{
                                background-color: #edb61f;
                                border: 1px solid black
                                }''')
        self.frame4.show()

    def frameMercadorias(self):
        self.frame5 = QFrame(self)
        self.frame5.setGeometry(0, 55, 850, 495)
        self.frame5.setStyleSheet("QFrame {background-color: #f0e1f5}")
        self.frame5.setFrameShape(QFrame.StyledPanel)
        self.frame5.setFrameShadow(QFrame.Raised)

        #textos
        font = QFont()
        font.setFamily("Helvética")
        font.setPointSize(16)
        font.setBold(True)
        font.setWeight(65)

        font2 = QFont()
        font2.setFamily('Helvética')
        font2.setPointSize(13)

        font3 = QFont()
        font3.setFamily('Calibri')
        font3.setPointSize(10)

        label = QLabel(self.frame5)
        label.setGeometry(40, 100, 71, 21)
        label.setText('Altura: ')
        label.setFont(font)

        label2 = QLabel(self.frame5)
        label2.setGeometry(440, 100, 180, 31)
        label2.setText('Quantidade:')
        label2.setFont(font)

        label3 = QLabel(self.frame5)
        label3.setGeometry(40, 220, 150, 31)
        label3.setText('Comprimento:')
        label3.setFont(font)

        label4 = QLabel(self.frame5)
        label4.setGeometry(40, 340, 161, 31)
        label4.setText('Largura:')
        label4.setFont(font)

        label5 = QLabel(self.frame5)
        label5.setGeometry(440, 220, 180, 31)
        label5.setText('Tipo de carga:')
        label5.setFont(font)

        label6 = QLabel(self.frame5)
        label6.setGeometry(222, 130, 60, 30)
        label6.setText('(cm)')
        label6.setFont(font2)

        label7 = QLabel(self.frame5)
        label7.setGeometry(222, 250, 60, 30)
        label7.setText('(cm)')
        label7.setFont(font2)

        label8 = QLabel(self.frame5)
        label8.setGeometry(222, 370, 60, 30)
        label8.setText('(cm)')
        label8.setFont(font2)

        label9 = QLabel(self.frame5)
        label9.setGeometry(490, 255, 180, 15)
        label9.setText('1-Cargas Frigorífico: alimentos')
        label9.setFont(font3)
        label9.setStyleSheet('color: gray')

        label10 = QLabel(self.frame5)
        label10.setGeometry(490, 275, 180, 15)
        label10.setText('2-Cargas Vivas: animais')
        label10.setFont(font3)
        label10.setStyleSheet('color: gray')

        label11 = QLabel(self.frame5)
        label11.setGeometry(490, 295, 350, 15)
        label11.setText('3-Cargas Secas: produtos em caixas (móveis, por exemplo)')
        label11.setFont(font3)
        label11.setStyleSheet('color: gray')

        label12 = QLabel(self.frame5)
        label12.setGeometry(490, 315, 350, 15)
        label12.setText('4-Cargas perigosas: explosivos ou vidro')
        label12.setFont(font3)
        label12.setStyleSheet('color: gray')

        #botões
        estilo_botao = '''QPushButton 
                                    {background-color: #9c32bf;
                                    border-radius: 10px;
                                    color: white;}
                                    QPushButton:hover {
                                        background-color: #fff;
                                        color: #9c32bf;
                                        border: 1px solid #9c32bf;
                                    }'''

        self.botao_limpar_mercadoria = QPushButton('Limpar', self.frame5)
        self.botao_limpar_mercadoria.setGeometry(440, 370, 80, 40)
        self.botao_limpar_mercadoria.setStyleSheet(estilo_botao)
        self.botao_limpar_mercadoria.setFont(QFont('Calibri', 15))
        self.botao_limpar_mercadoria.clicked.connect(self.limparMercadorias)

        self.botao_enviar_mercadoria = QPushButton('Enviar', self.frame5)
        self.botao_enviar_mercadoria.setGeometry(580, 370, 80, 40)
        self.botao_enviar_mercadoria.setStyleSheet(estilo_botao)
        self.botao_enviar_mercadoria.setFont(QFont('Calibri', 15))
        self.botao_enviar_mercadoria.clicked.connect(self.frameLocalizacao)

        #entradas
        estilo = '''QLineEdit{background-color: white;
                                    padding-left: 10px;
                                    font-size: 11pt; 
                                    font-family: arial;
                                    font-weight: 480;
                                    color: rgba(0, 0, 0, 0.8)}''' 

        self.entrada_altura = QLineEdit(self.frame5)
        self.entrada_altura.setGeometry(40, 130, 181, 31)
        self.entrada_altura.setStyleSheet(estilo)

        self.entrada_quantidade = QLineEdit(self.frame5)
        self.entrada_quantidade.setGeometry(440, 130, 130, 31)
        self.entrada_quantidade.setStyleSheet(estilo)

        self.entrada_comprimento = QLineEdit(self.frame5)
        self.entrada_comprimento.setGeometry(40, 250, 181, 31)
        self.entrada_comprimento.setStyleSheet(estilo)

        self.entrada_carga = QLineEdit(self.frame5)
        self.entrada_carga.setGeometry(440, 250, 40, 31)
        self.entrada_carga.setStyleSheet(estilo)

        self.entrada_largura = QLineEdit(self.frame5)
        self.entrada_largura.setGeometry(40, 370, 181, 31)
        self.entrada_largura.setStyleSheet(estilo)
        self.frame5.show()

    def limparMercadorias(self):
        self.entrada_altura.clear()
        self.entrada_quantidade.clear()
        self.entrada_comprimento.clear()
        self.entrada_carga.clear()
        self.entrada_largura.clear()

    def frameLocalizacao(self):
        self.altura = self.entrada_altura.text()
        self.quantidade = self.entrada_quantidade.text()
        self.comprimento = self.entrada_comprimento.text()
        self.tipo_de_carga = self.entrada_carga.text()
        self.largura = self.entrada_largura.text()

        self.conectar_no_banco()
        self.cursor.execute(f'''SELECT id_cliente FROM clientes WHERE cpf = '{self.cpf}'; ''')
        self.id_user = self.cursor.fetchall()
        self.cursor.execute(f'''INSERT INTO mercadorias(id_fk_cliente, altura, comprimento, largura, quantidade, tipo_de_carga)
                            VALUES('{self.id_user[0][0]}', '{self.altura}', '{self.comprimento}', '{self.largura}', '{self.quantidade}', '{self.tipo_de_carga}');''')
        
        sip.delete(self.frame5)

        self.frame6 = QFrame(self)
        self.frame6.setGeometry(0, 55, 850, 495)
        self.frame6.setStyleSheet("QFrame {background-color: #f0e1f5}")
        self.frame6.setFrameShape(QFrame.StyledPanel)
        self.frame6.setFrameShadow(QFrame.Raised)

        #textos
        font = QFont()
        font.setFamily('Helvética')
        font.setPointSize(16)
        font.setBold(True)
        font.setWeight(65)

        font2 = QFont()
        font2.setFamily('Helvética')
        font2.setPointSize(14)
        font2.setBold(True)
        font2.setWeight(65)

        font3 = QFont()
        font3.setFamily('Calibri')
        font3.setPointSize(16)

        label = QLabel(self.frame6)
        label.setGeometry(40, 30, 80, 30)
        label.setText('Origem:')
        label.setFont(font)

        label2 = QLabel(self.frame6)
        label2.setGeometry(40, 180, 90, 30)
        label2.setText('Destino:')
        label2.setFont(font)

        label3 = QLabel(self.frame6)
        label3.setGeometry(40, 70, 700, 40)
        label3.setText('  R. Guanabara, 217 - St. Pausanes, Rio Verde - GO, 75904-015')
        label3.setStyleSheet('''background-color: #dbdad7;
                                border: 0.5px solid;
                                border-top-color: #dbdad7;
                                border-left-color: #dbdad7;
                                border-right-color: #dbdad7;
                                border-bottom-color: black''')
        label3.setFont(font3)

        label4 = QLabel(self.frame6)
        label4.setGeometry(40, 290, 100, 30)
        label4.setText('Bairro:')
        label4.setFont(font2)

        label5 = QLabel(self.frame6)
        label5.setGeometry(350, 290, 100, 30)
        label5.setText('Rua:')
        label5.setFont(font2)

        label6 = QLabel(self.frame6)
        label6.setGeometry(40, 390, 100, 30)
        label6.setText('Quadra:')
        label6.setFont(font2)

        label7 = QLabel(self.frame6)
        label7.setGeometry(350, 390, 100, 30)
        label7.setText('Lote:')
        label7.setFont(font2)

        self.label_vazia_quadra = QLabel(self.frame6)
        self.label_vazia_quadra.setGeometry(40, 320, 200, 30)
        self.label_vazia_quadra.setStyleSheet('background-color: #fff')

        self.label_vazia_lote = QLabel(self.frame6)
        self.label_vazia_lote.setGeometry(350, 320, 350, 30)
        self.label_vazia_lote.setStyleSheet('background-color: #fff')

        #botões
        estilo_botao = '''QPushButton 
                                    {background-color: #9c32bf;
                                    border-radius: 10px;
                                    color: white;}
                                    QPushButton:hover {
                                        background-color: #fff;
                                        color: #9c32bf;
                                        border: 1px solid #9c32bf;
                                    }'''

        self.botao_limpar_localizacao = QPushButton('Limpar', self.frame6)
        self.botao_limpar_localizacao.setGeometry(620, 420, 80, 40)
        self.botao_limpar_localizacao.setStyleSheet(estilo_botao)
        self.botao_limpar_localizacao.setFont(QFont('Calibri', 15))
        self.botao_limpar_localizacao.clicked.connect(self.limparLocalizacao)

        self.botao_enviar_localizacao = QPushButton('Enviar', self.frame6)
        self.botao_enviar_localizacao.setGeometry(720, 420, 80, 40)
        self.botao_enviar_localizacao.setStyleSheet(estilo_botao)
        self.botao_enviar_localizacao.setFont(QFont('Calibri', 15))
        self.botao_enviar_localizacao.clicked.connect(self.framePreco)

        self.botao_ok = QPushButton('Ok', self.frame6)
        self.botao_ok.setGeometry(720, 220, 40, 40)
        self.botao_ok.setStyleSheet(estilo_botao)
        self.botao_ok.setFont(QFont('Calibri', 15))
        self.botao_ok.clicked.connect(self.informacoes_cep)

        #entradas
        estilo = '''QLineEdit{background-color: white;
                                    padding-left: 10px;
                                    font-size: 11pt; 
                                    font-family: arial;
                                    font-weight: 480;
                                    color: rgba(0, 0, 0, 0.8)}'''
        estilo2 = '''QLineEdit{background-color: #fff;
                                    padding-left: 10px;
                                    font-size: 11pt; 
                                    font-family: arial;
                                    font-weight: 480;
                                    color: gray}''' 

        self.entrada_cep = QLineEdit(self.frame6)
        self.entrada_cep.setGeometry(40, 220, 630, 40)
        self.entrada_cep.setStyleSheet(estilo)
        self.entrada_cep.setPlaceholderText('Digite o CEP do local de origem:')

        self.entrada_quadra = QLineEdit(self.frame6)
        self.entrada_quadra.setGeometry(40, 421, 170, 30)
        self.entrada_quadra.setStyleSheet(estilo2)

        self.entrada_lote = QLineEdit(self.frame6)
        self.entrada_lote.setGeometry(350, 421, 170, 30)
        self.entrada_lote.setStyleSheet(estilo2)
        self.frame6.show()

    def limparLocalizacao(self):
        self.entrada_cep.clear()
        self.entrada_quadra.clear()
        self.entrada_lote.clear()

    def informacoes_cep(self):
        self.cep = self.entrada_cep.text()
        self.label_vazia_quadra.close()
        self.label_vazia_lote.close()

        requerimento = requests.get(f'https://viacep.com.br/ws/{self.cep}/json/')
        informacoes = requerimento.json()
        self.bairro = informacoes['bairro']
        self.rua = informacoes['logradouro']

        font4 = QFont()
        font4.setFamily('Calibri')
        font4.setPointSize(11)
        font4.setBold(True)

        label12 = QLabel(self.frame6)
        label12.setGeometry(40, 320, 200, 30)
        label12.setText(f'{self.bairro}')
        label12.setFont(font4)
        label12.setStyleSheet('background-color: #fff; color: gray')
        label12.show()

        label13 = QLabel(self.frame6)
        label13.setGeometry(350, 320, 350, 30)
        label13.setText(f'{self.rua}')
        label13.setFont(font4)
        label13.setStyleSheet('background-color: #fff; color: gray')
        label13.show()

        return self.bairro, self.rua

    def framePreco(self):
        #calcular distância e tempo entre 2 pontos com Api Matrix do googlemaps
        try:
            gmaps = googlemaps.Client(key='AIzaSyBgvdOnxyyeudB5VWkdjriDLDLMZNQODmQ')
            consulta = gmaps.distance_matrix('75904015', f'{self.cep}')
            self.distancia = consulta['rows'][0]['elements'][0]['distance']['text']
            distanciamenor = float(self.distancia[:-3])
            self.tempo = consulta['rows'][0]['elements'][0]['duration']['text']
            print(self.distancia, self.tempo)
        except:
            print('Seu CEP está incorreto ou não existe.')

        #calcular valor do frete
        km_rodado = 2.75 * distanciamenor
        volume = (float(self.altura) * float(self.comprimento) * float(self.largura))*0.0029
        fixo_volume = volume * int(self.quantidade)
        self.valor_pagar = km_rodado + fixo_volume
        print(f'{self.valor_pagar}')

        self.cep = self.cep
        self.bairro = self.bairro
        self.rua = self.rua
        self.quadra = self.entrada_quadra.text()
        self.lote = self.entrada_lote.text()

        self.conectar_no_banco()
        self.cursor.execute(f'''INSERT INTO localizacao(fk_id_cliente, cep, bairro, rua, quadra, lote)
                            VALUES('{self.id_user[0][0]}', '{self.cep}', '{self.bairro}', '{self.rua}', '{self.quadra}', '{self.lote}');''')

        #informações na tabela caminhao
        self.cursor.execute(f'''SELECT id_mercadoria FROM mercadorias WHERE id_fk_cliente = '{self.id_user[0][0]}'; ''')
        id_mercadory = self.cursor.fetchall()
        self.cursor.execute(f'''INSERT INTO caminhao(fk_cep_cliente, fk_mercadoria, cep_caminhao)
                            VALUES('{self.cep}', '{id_mercadory[0][0]}', '75904015'); ''')

        sip.delete(self.frame6)
        sip.delete(self.frame4)

        self.setGeometry(450, 110, 500, 500)
        self.setWindowTitle('TransportFlash')
        self.setStyleSheet('background-color: #f0e1f5')
        self.icone = 'minicaminhao.png'
        self.setWindowIcon(QtGui.QIcon(self.icone))

        self.frame7 = QFrame(self)
        self.frame7.setGeometry(0, 0, 500, 500)
        self.frame7.setStyleSheet("QFrame {background-color: #f0e1f5}")
        self.frame7.setFrameShape(QFrame.StyledPanel)
        self.frame7.setFrameShadow(QFrame.Raised)

        font = QFont()
        font.setFamily("Cascadia Mono SemiBold")
        font.setPointSize(20)
        font2 = QFont()
        font2.setFamily("Cascadia Mono SemiBold")
        font2.setPointSize(14)

        #textos
        label = QLabel(self.frame7)
        label.setText('Valor:')
        label.setGeometry(130, 50, 300, 30)
        label.setFont(font)

        label2 = QLabel(self.frame7)
        label2.setText(f'R${self.valor_pagar:.2f}')
        label2.setGeometry(230, 50, 300, 30)
        label2.setStyleSheet('color: red')
        label2.setFont(font)

        label3 = QLabel(self.frame7)
        label3.setText('Tempo:')
        label3.setGeometry(130, 100, 300, 30)
        label3.setFont(font)

        label4 = QLabel(self.frame7)
        label4.setText(f'{self.tempo}')
        label4.setGeometry(230, 100, 300, 30)
        label4.setStyleSheet('color: red')
        label4.setFont(font)

        label5 = QLabel(self.frame7)
        label5.setText('Escolha sua forma de pagamento: ')
        label5.setGeometry(20, 160, 500, 40)
        label5.setFont(font2)

        #botões
        estilo_botao = '''QPushButton 
                            {background-color: #9c32bf;
                            border-radius: 10px;
                            color: white;}
                            QPushButton:hover {
                                background-color: #fff;
                                color: #9c32bf;
                                border: 1px solid #9c32bf;
                            }'''

        self.botao_cancelar = QPushButton('Cancelar', self.frame7)
        self.botao_cancelar.setGeometry(285, 430, 90, 35)
        self.botao_cancelar.setStyleSheet(estilo_botao)
        self.botao_cancelar.setFont(QFont('Calibri', 15))

        self.botao_confirmar = QPushButton('Confirmar', self.frame7)
        self.botao_confirmar.setGeometry(385, 430, 90, 35)
        self.botao_confirmar.setStyleSheet(estilo_botao)
        self.botao_confirmar.setFont(QFont('Calibri', 15))
        self.botao_confirmar.clicked.connect(self.frameConfirmacao)

        self.botao_credito = QPushButton('''Cartão de\nCrédito''', self.frame7)
        self.botao_credito.setGeometry(20, 210, 90, 50)
        self.botao_credito.setStyleSheet(estilo_botao)
        fontbutao = QFont()
        fontbutao.setFamily('Calibri')
        fontbutao.setPointSize(13)
        fontbutao.setBold(True)
        self.botao_credito.setFont(fontbutao)

        self.botao_debito = QPushButton('''Cartão de\ndébito''', self.frame7)
        self.botao_debito.setGeometry(140, 210, 90, 50)
        self.botao_debito.setStyleSheet(estilo_botao)
        fontbutao = QFont()
        fontbutao.setFamily('Calibri')
        fontbutao.setPointSize(13)
        fontbutao.setBold(True)
        self.botao_debito.setFont(fontbutao)

        self.botao_pix = QPushButton('Pix', self.frame7)
        self.botao_pix.setGeometry(260, 210, 90, 50)
        self.botao_pix.setStyleSheet(estilo_botao)
        fontbutao = QFont()
        fontbutao.setFamily('Calibri')
        fontbutao.setPointSize(15)
        fontbutao.setBold(True)
        self.botao_pix.setFont(fontbutao)

        self.botao_real = QPushButton('Dinheiro', self.frame7)
        self.botao_real.setGeometry(380, 210, 90, 50)
        self.botao_real.setStyleSheet(estilo_botao)
        fontbutao = QFont()
        fontbutao.setFamily('Calibri')
        fontbutao.setPointSize(15)
        fontbutao.setBold(True)
        self.botao_real.setFont(fontbutao)

        #entrada
        self.entrada_pagar = QLineEdit(self.frame7)
        self.entrada_pagar.setGeometry(70, 310, 350, 40)
        self.entrada_pagar.setStyleSheet('background-color: #fff')
        self.frame7.show()
        
    def frameConfirmacao(self):
        sip.delete(self.frame7)

        self.setGeometry(450, 250, 500, 200)
        self.setWindowTitle('TransportFlash')
        self.setStyleSheet('background-color: #f0e1f5')
        self.icone = 'minicaminhao.png'
        self.setWindowIcon(QtGui.QIcon(self.icone))

        self.frame8 = QFrame(self)
        self.frame8.setGeometry(0, 0, 500, 200)
        self.frame8.setStyleSheet("QFrame {background-color: #f0e1f5}")
        self.frame8.setFrameShape(QFrame.StyledPanel)
        self.frame8.setFrameShadow(QFrame.Raised)

        #texto
        font = QFont()
        font.setFamily("Helvética")
        font.setPointSize(20)
        label = QLabel(self.frame8)
        label.setText('Compra finalizada!')
        label.setGeometry(130, 70, 300, 30)
        label.setStyleSheet('color: #9c32bf')
        label.setFont(font)

        #botoes
        estilo_botao = '''QPushButton 
                            {background-color: #9c32bf;
                            border-radius: 10px;
                            color: white;}
                            QPushButton:hover {
                                background-color: #fff;
                                color: #9c32bf;
                                border: 1px solid #9c32bf;
                            }'''

        self.botao_yes = QPushButton('Ok', self.frame8)
        self.botao_yes.setGeometry(400, 140, 90, 35)
        self.botao_yes.setStyleSheet(estilo_botao)
        self.botao_yes.setFont(QFont('Calibri', 15))
        self.botao_yes.clicked.connect(self.frameBoleto)
        self.frame8.show()

    def frameBoleto(self):
        sip.delete(self.frame8)
        self.setGeometry(450, 250, 500, 200)
        self.setWindowTitle('TransportFlash')
        self.setStyleSheet('background-color: #f0e1f5')
        self.icone = 'minicaminhao.png'
        self.setWindowIcon(QtGui.QIcon(self.icone))

        self.frame9 = QFrame(self)
        self.frame9.setGeometry(0, 0, 500, 200)
        self.frame9.setStyleSheet("QFrame {background-color: #f0e1f5}")
        self.frame9.setFrameShape(QFrame.StyledPanel)
        self.frame9.setFrameShadow(QFrame.Raised)

        #texto
        font = QFont()
        font.setFamily("Helvética")
        font.setPointSize(13)
        label = QLabel(self.frame9)
        label.setText('   Baixe o pdf do seu boleto clicando no botão abaixo.\nEm seguida imprima-o e pague na lotérica mais próxima.')
        label.setGeometry(40, 50, 500, 70)
        label.setStyleSheet('color: #9c32bf')
        label.setFont(font)

        #botoes
        self.botao1 = QPushButton('Baixar', self.frame9)
        self.botao1.setGeometry(210, 140, 90, 35)
        self.botao1.setStyleSheet('''QPushButton 
                                    {background-color: #9c32bf;
                                    border-radius: 10px;
                                    color: white;}
                                    QPushButton:hover {
                                        background-color: #fff;
                                        color: #9c32bf;
                                        border: 1px solid #9c32bf;
                                    }''')
        self.botao1.setFont(QFont('Calibri', 15))
        self.botao1.clicked.connect(self.baixa_boleto)
        self.frame9.show()

    def baixa_boleto(self):
        self.loli = self.id_user
        self.conectar_no_banco()
        self.cursor.execute(f'''SELECT nome FROM clientes WHERE id_cliente = '{self.loli[0][0]}'; ''')
        nome_usuario = self.cursor.fetchall()
        print(nome_usuario)

        data_hoje = datetime.today().strftime('%d-%m-%Y')
        data_em_banco = datetime.today().strftime('%Y/%m/%d')
        numero_carga = str(self.tipo_de_carga)
        if numero_carga == '1':
            carga = 'Carga Frigorífica'
        elif numero_carga == '2':
            carga = 'Cargas Vivas'
        elif numero_carga == '3':
            carga = 'Cargas Secas'
        else:
            carga = 'Cargas Perigosas'

        #informacoes nas tabelas de boleto
        self.cursor.execute(f'''INSERT INTO boletos(id_fk_cliente, valor, data, fk_nome_cliente, fk_quantidade, fk_tipo_carga, fk_cep, fk_bairro, fk_rua, fk_quadra, fk_lote)
        VALUES('{self.id_user[0][0]}', '{self.valor_pagar:.2f}', '{data_em_banco}', '{str(nome_usuario[0][0])}', '{self.quantidade}', '{self.tipo_de_carga}', '{self.cep}', '{self.bairro}', '{self.rua}', '{self.quadra}', '{self.lote}'); ''')

        documento = Document()
        linha1 = documento.add_paragraph()
        linha1.add_run('Termos e condições de adiamento contratual \nem célula de crédito bancário').bold = True
        linha1_format = linha1.paragraph_format
        linha1_format.alignment
        linha1_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        linha2 = documento.add_paragraph().add_run(f'Prezado(a) {nome_usuario[0][0]},\nConforme condições abaixo, o pagamento deste boleto refere-se à parcela de entrada da renegociação da CCB/Contrato n.10-427359/15, firmado junto ao TransportFlash.')
        fonte = linha2.font
        fonte.name = 'Calibri'
        fonte.size = Pt(13)

        linha3 = documento.add_paragraph().add_run('DISTRIBUIDORA TRANSPORTFLASH DE GOIÁS    CNPJ: 03.786.187/0018-37,\nRUA GUANABARA 237, SETOR PAUZANES - Rio Verde/GO - CEP. 75.904-015')
        fonte = linha3.font

        linha_vazia = documento.add_paragraph('')
        linha2_vazia = documento.add_paragraph('')
        linha3_vazia = documento.add_paragraph('')

        linha4 = documento.add_paragraph().add_run('TRANSPORTFLASH                       23791,11103 55844.022568')
        fonte2 = linha4.font
        fonte2.size = Pt(15)
        fonte2.bold = True

        data = ( 
                ('Local de pagamento:\nPAGÁVEL PREFERENCIALMENTE NAS LOTÉRICAS', f'Data do documento:\n{data_hoje}'), 
                (f'Quantidade:\n{self.quantidade}', f'Valor do documento:\nR${self.valor_pagar:.2f}')
            ) 
        table = documento.add_table(rows = 2, cols = 2) 
        row = table.rows[0].cells 

        for col1, col2 in data:
            row = table.add_row().cells     
            row[0].text = col1
            row[1].text = col2 

        linha5 = documento.add_paragraph().add_run(f'Instruções (Texto de responsabilidade do cedente)\nTIPO DE CARGA: \n {carga}')
        fonte3 = linha5.font
        fonte3.name = 'Calibri'
        fonte3.size = Pt(13)

        linha4_vazia = documento.add_paragraph('')
        linha5_vazia = documento.add_paragraph('')
        linha6_vazia = documento.add_paragraph('')

        linha6 = documento.add_paragraph().add_run(f'Pagador\nBairro {self.bairro}, Rua {self.rua}\nQuadra {self.quadra}, Lote{self.lote}, CEP: {self.cep}')

        documento.save('Meu boleto.docx')

        wdFormatPDF = 17
        inputFile = os.path.abspath("Meu boleto.docx")
        outputFile = os.path.abspath("Meu_boleto.pdf")
        word = win32com.client.Dispatch('Word.Application')
        doc = word.Documents.Open(inputFile)
        doc.SaveAs(outputFile, FileFormat=wdFormatPDF)
        doc.Close()
        word.Quit()

aplicacao = QApplication(sys.argv)
x = JanelaInicial()
x.show()
sys.exit(aplicacao.exec_())